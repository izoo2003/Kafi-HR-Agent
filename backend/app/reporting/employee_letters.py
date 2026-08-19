"""Fill official appointment / contract Word templates from the employee record."""
from __future__ import annotations

import re
from datetime import date, datetime
from decimal import Decimal
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.core.config import get_settings
from app.models.employees import Employee, EmployeeReference

_TOKEN = re.compile(r"\{\{\s*([a-z0-9_]+)\s*\}\}")

_KIND_FILES = {
    "appointment": "appointment_letter.docx",
    "contract": "employment_contract.docx",
}

_EMPLOYMENT_LABELS = {
    "full_time": "full-time",
    "part_time": "part-time",
    "contract": "contract",
}

_BLOOD_HINTS = (
    "mother",
    "father",
    "brother",
    "sister",
    "son",
    "daughter",
    "wife",
    "husband",
    "spouse",
    "uncle",
    "aunt",
    "cousin",
    "blood",
    "parent",
    "relative",
)

_ONES = [
    "",
    "One",
    "Two",
    "Three",
    "Four",
    "Five",
    "Six",
    "Seven",
    "Eight",
    "Nine",
    "Ten",
    "Eleven",
    "Twelve",
    "Thirteen",
    "Fourteen",
    "Fifteen",
    "Sixteen",
    "Seventeen",
    "Eighteen",
    "Nineteen",
]
_TENS = ["", "", "Twenty", "Thirty", "Forty", "Fifty", "Sixty", "Seventy", "Eighty", "Ninety"]


def letter_kinds() -> tuple[str, ...]:
    return tuple(_KIND_FILES.keys())


def _fmt_date(value: date | datetime | None, fmt: str) -> str:
    if value is None:
        return "[date to be confirmed]"
    if isinstance(value, datetime):
        value = value.date()
    return value.strftime(fmt)


def _int_salary(value: Decimal | float | int | None) -> int | None:
    if value is None:
        return None
    return int(Decimal(str(value)).quantize(Decimal("1")))


def _fmt_amount(n: int | None) -> str:
    if n is None:
        return "[to be confirmed]"
    return f"{n:,}"


def _under_thousand(n: int) -> str:
    if n < 20:
        return _ONES[n]
    if n < 100:
        tens, ones = divmod(n, 10)
        return f"{_TENS[tens]} {_ONES[ones]}".strip()
    hundreds, rest = divmod(n, 100)
    if rest == 0:
        return f"{_ONES[hundreds]} Hundred"
    return f"{_ONES[hundreds]} Hundred {_under_thousand(rest)}"


def rupees_in_words(amount: int | None) -> str:
    if amount is None:
        return "amount to be confirmed"
    if amount == 0:
        return "Zero"
    n = amount
    parts: list[str] = []
    crore, n = divmod(n, 10_000_000)
    lakh, n = divmod(n, 100_000)
    thousand, n = divmod(n, 1_000)
    if crore:
        parts.append(f"{_under_thousand(crore)} Crore")
    if lakh:
        parts.append(f"{_under_thousand(lakh)} Lakh")
    if thousand:
        parts.append(f"{_under_thousand(thousand)} Thousand")
    if n:
        parts.append(_under_thousand(n))
    return " ".join(parts)


def _salary_breakup(total: int | None) -> tuple[int | None, int | None, int | None]:
    """Keep the official 120 : 45 : 30 split of the sample package (195 total)."""
    if total is None:
        return None, None, None
    basic = round(total * 120 / 195)
    hra = round(total * 45 / 195)
    allowance = total - basic - hra
    return basic, hra, allowance


def _honorific(gender: str | None) -> str:
    g = (gender or "").strip().lower()
    if g in {"female", "f", "woman", "ms", "mrs", "miss"}:
        return "Ms."
    return "Mr."


def _filiation(gender: str | None) -> str:
    g = (gender or "").strip().lower()
    if g in {"female", "f", "woman", "ms", "mrs", "miss"}:
        return "D/o."
    return "S/o."


def _a_or_an(phrase: str) -> str:
    first_word = (phrase or "").strip().split(" ", 1)[0]
    if first_word.isupper() and 1 <= len(first_word) <= 4:
        letter = first_word[0].lower()
        return "an" if letter in "aefhilmnorsx" else "a"
    first = re.sub(r"[^A-Za-z]", "", first_word)[:1].lower()
    return "an" if first in {"a", "e", "i", "o", "u"} else "a"


def _is_blood_relative(ref: EmployeeReference) -> bool:
    blob = f"{ref.relation or ''} {ref.notes or ''}".lower()
    return any(hint in blob for hint in _BLOOD_HINTS)


def _format_reference(ref: EmployeeReference | None, index: int) -> str:
    if ref is None:
        return f"{index}. ________________ (________________) = (________________)"
    title = (ref.relation or "").strip() or "Reference"
    phone = (ref.phone or "").strip() or "________________"
    return f"{index}. {ref.full_name} ({title}) = ({phone})"


def _reference_lines(employee: Employee) -> dict[str, str]:
    refs = list(employee.references or [])
    blood = [r for r in refs if _is_blood_relative(r)]
    professional: list[EmployeeReference | None] = [r for r in refs if r not in blood]
    while len(professional) < 3:
        professional.append(None)
    blood_one = blood[0] if blood else None
    return {
        "ref_pro_1": _format_reference(professional[0], 1),
        "ref_pro_2": _format_reference(professional[1], 2),
        "ref_pro_3": _format_reference(professional[2], 3),
        "ref_blood_1": (
            f"1.{blood_one.full_name} ({blood_one.relation}) = ({blood_one.phone or '________________'})"
            if blood_one
            else "1.________________ (________________) = (________________)"
        ),
    }


def merge_fields(employee: Employee) -> dict[str, str]:
    dept = getattr(employee, "department", None)
    dept_name = getattr(dept, "name", None) if dept is not None else None
    emp_type = (employee.employment_type or "full_time").strip()
    total = _int_salary(employee.base_salary)
    basic, hra, allowance = _salary_breakup(total)
    role = (employee.role_title or dept_name or "Employee").strip()
    honorific = _honorific(employee.gender)
    father = (employee.father_name or "[father's name]").strip()
    if father.lower().startswith("mr."):
        father_honorific = "Mr."
        father = father[3:].strip()
    else:
        father_honorific = "Mr."
    phone = (employee.personal_mobile or employee.alternate_mobile or "").strip()
    email = (employee.email or "").strip()
    phone_email = " ".join(p for p in (phone, email) if p) or "[to be confirmed]"
    today = date.today()
    joined = employee.date_joined
    fields = {
        "company_name": "Kafi Group",
        "full_name": (employee.full_name or "").strip(),
        "honorific": honorific,
        "filiation": _filiation(employee.gender),
        "father_honorific": father_honorific,
        "father_name": father,
        "cnic": (employee.cnic or "[CNIC]").strip(),
        "employee_code": employee.employee_code or "",
        "role_title": role,
        "a_or_an": _a_or_an(role),
        "department": dept_name or "[department]",
        "employment_type": _EMPLOYMENT_LABELS.get(emp_type, emp_type.replace("_", " ")),
        "dated": _fmt_date(today, "%B %d, %Y"),
        "today": _fmt_date(today, "%d %B %Y"),
        "date_joined": _fmt_date(joined, "%d %B %Y"),
        "date_joined_long": _fmt_date(joined, "%d-%B-%Y"),
        "date_joined_contract": _fmt_date(joined, "%d-%B %Y"),
        "date_joined_short": _fmt_date(joined, "%d-%b-%Y"),
        "base_salary": _fmt_amount(total),
        "salary_rs": _fmt_amount(total),
        "salary_basic": _fmt_amount(basic),
        "salary_hra": _fmt_amount(hra),
        "salary_allowance": _fmt_amount(allowance),
        "salary_words": rupees_in_words(total),
        "base_salary_words_hint": (
            "Pakistan Rupees" if total is not None else "salary to be confirmed"
        ),
        "email": email or "—",
        "mobile": phone or "—",
        "phone_and_email": phone_email,
        "city": employee.city or "",
        "current_address": employee.current_address or employee.permanent_address or "—",
        "permanent_address": employee.permanent_address or employee.current_address or "—",
        "nationality": employee.nationality or "Pakistani",
        "gender": employee.gender or "—",
    }
    fields.update(_reference_lines(employee))
    return fields


def _fill_text(text: str, fields: dict[str, str]) -> str:
    def repl(match: re.Match[str]) -> str:
        key = match.group(1)
        return fields.get(key, match.group(0))

    return _TOKEN.sub(repl, text)


def _template_path(kind: str) -> Path:
    filename = _KIND_FILES[kind]
    bundled = Path(__file__).resolve().parents[2] / "config" / "letter_templates" / filename
    configured = get_settings().config_dir / "letter_templates" / filename
    for path in (configured, bundled):
        if path.is_file():
            return path
    return bundled


def _set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(text)
        return
    runs[0].text = text
    for run in runs[1:]:
        run.text = ""


def _walk_tables(tables: list[Table]):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
                yield from _walk_tables(cell.tables)


def _iter_paragraphs(doc: DocumentObject):
    yield from doc.paragraphs
    yield from _walk_tables(doc.tables)
    for section in doc.sections:
        for attr in ("header", "footer", "first_page_header", "first_page_footer"):
            try:
                part = getattr(section, attr)
            except Exception:
                continue
            yield from part.paragraphs
            yield from _walk_tables(part.tables)


def _replace_xml_text_nodes(doc: DocumentObject, fields: dict[str, str]) -> None:
    """Catch leftover tokens in text boxes / SDTs that python-docx paragraphs miss."""
    for node in doc.element.iter(qn("w:t")):
        if node.text and "{{" in node.text:
            node.text = _fill_text(node.text, fields)


def fill_document(doc: DocumentObject, fields: dict[str, str]) -> None:
    for paragraph in _iter_paragraphs(doc):
        original = paragraph.text
        if "{{" not in original:
            continue
        _set_paragraph_text(paragraph, _fill_text(original, fields))
    _replace_xml_text_nodes(doc, fields)


def _slug(name: str) -> str:
    cleaned = re.sub(r"[^\w\-]+", "_", (name or "").strip(), flags=re.UNICODE)
    return (cleaned.strip("_") or "employee")[:60]


def render_docx_bytes(kind: str, employee: Employee) -> tuple[bytes, str]:
    if kind not in _KIND_FILES:
        raise ValueError(f"Unknown letter type '{kind}'")
    path = _template_path(kind)
    if not path.is_file():
        raise FileNotFoundError(f"Letter template missing: {path}")
    fields = merge_fields(employee)
    doc = Document(str(path))
    fill_document(doc, fields)
    buf = BytesIO()
    doc.save(buf)
    title_slug = "Appointment_Letter" if kind == "appointment" else "Employment_Contract"
    filename = f"{title_slug}_{_slug(employee.full_name)}.docx"
    return buf.getvalue(), filename
