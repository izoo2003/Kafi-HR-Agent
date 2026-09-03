"""Generate appointment letters and employment contracts for an employee."""
from __future__ import annotations

import json
import logging
import re
from io import BytesIO
from typing import Any

from sqlalchemy.orm import Session

from app.core.config import get_settings
from app.core.exceptions import BusinessRuleViolation, EntityNotFound, ValidationFailed
from app.core.gemini_client import generate_content_with_fallback
from app.ingestion.employee_docs import delete_stored_file, read_stored_file, store_employee_file, store_letter_file
from app.models.employees import EmployeeDocument
from app.reporting.employee_letters import letter_kinds, render_docx_bytes
from app.schemas.common import AuthContext
from app.schemas.employees import LETTER_CATEGORIES, LETTER_SIGNED_CATEGORIES, LetterSignatureVerifyResult
from app.services import audit_service, employee_service

logger = logging.getLogger(__name__)

DOCX_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
VERIFY_IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".webp", ".gif", ".heic", ".heif"}
VERIFY_PDF_SUFFIXES = {".pdf"}

_MISSING = "It is not created yet. Create them first."


def _verify_upload_mime(filename: str, mime_type: str | None) -> str:
    """Return a mime type Gemini can read for signature verification."""
    suffix = ""
    if "." in (filename or ""):
        suffix = "." + filename.rsplit(".", 1)[-1].lower()
    mime = (mime_type or "").lower().strip()
    if mime.startswith("image/") or suffix in VERIFY_IMAGE_SUFFIXES:
        if mime.startswith("image/"):
            return mime
        guessed = {
            ".png": "image/png",
            ".jpg": "image/jpeg",
            ".jpeg": "image/jpeg",
            ".webp": "image/webp",
            ".gif": "image/gif",
            ".heic": "image/heic",
            ".heif": "image/heif",
        }.get(suffix, "image/jpeg")
        return guessed
    if mime == "application/pdf" or suffix in VERIFY_PDF_SUFFIXES:
        return "application/pdf"
    raise ValidationFailed(
        "Upload a PDF or image of the signed letter (PDF, PNG, or JPG)."
    )


def _category(kind: str) -> str:
    if kind not in LETTER_CATEGORIES:
        raise ValidationFailed(f"Unknown letter type '{kind}'")
    return LETTER_CATEGORIES[kind]


def _signed_category(kind: str) -> str:
    if kind not in LETTER_SIGNED_CATEGORIES:
        raise ValidationFailed(f"Unknown letter type '{kind}'")
    return LETTER_SIGNED_CATEGORIES[kind]


def _latest_letter(db: Session, employee_id: int, kind: str) -> EmployeeDocument | None:
    cat = _category(kind)
    return (
        db.query(EmployeeDocument)
        .filter(EmployeeDocument.employee_id == employee_id, EmployeeDocument.category == cat)
        .order_by(EmployeeDocument.id.desc())
        .first()
    )


def _delete_category_docs(db: Session, employee_id: int, category: str) -> None:
    existing = (
        db.query(EmployeeDocument)
        .filter(EmployeeDocument.employee_id == employee_id, EmployeeDocument.category == category)
        .all()
    )
    for doc in existing:
        path = doc.file_path
        db.delete(doc)
        db.flush()
        delete_stored_file(path)


def generate_letter(
    db: Session,
    auth: AuthContext,
    employee_id: int,
    kind: str,
) -> tuple[bytes, str]:
    employee = employee_service.get_employee(db, employee_id)
    try:
        content, filename = render_docx_bytes(kind, employee)
    except FileNotFoundError as exc:
        raise ValidationFailed(str(exc)) from exc
    except Exception as exc:
        logger.exception("Letter render failed for employee %s kind %s", employee_id, kind)
        raise ValidationFailed(f"Could not generate letter: {exc}") from exc
    cat = _category(kind)
    # Recreate clears prior letter + any signed/verified scan
    _delete_category_docs(db, employee.id, cat)
    _delete_category_docs(db, employee.id, _signed_category(kind))

    stored_path, mime = store_letter_file(
        employee_id=employee.id,
        filename=filename,
        content=content,
    )
    row = EmployeeDocument(
        employee_id=employee.id,
        category=cat,
        title="Appointment letter" if kind == "appointment" else "Employment contract",
        file_path=stored_path,
        original_filename=filename,
        mime_type=mime or DOCX_TYPE,
    )
    db.add(row)
    db.flush()
    audit_service.log_from_auth(
        db,
        auth,
        action=f"employee.letter_{kind}",
        entity_type="employee",
        entity_id=employee.id,
        after_state={"filename": filename, "kind": kind, "document_id": row.id},
    )
    return content, filename


def view_stored_letter(db: Session, employee_id: int, kind: str) -> tuple[bytes, str]:
    employee_service.get_employee(db, employee_id)
    if kind not in letter_kinds():
        raise ValidationFailed(f"Unknown letter type '{kind}'")
    doc = _latest_letter(db, employee_id, kind)
    if doc is None:
        raise EntityNotFound(_MISSING)
    return read_stored_file(doc.file_path), doc.original_filename or f"{kind}.docx"


_APPOINTMENT_KINDS = frozenset(
    {
        "appointment",
        "appointment_letter",
        "appointment letter",
        "offer_letter",
        "offer letter",
        "letter of appointment",
    }
)
_CONTRACT_KINDS = frozenset(
    {
        "contract",
        "employment_contract",
        "employment contract",
        "contract_letter",
        "contract letter",
        "contract of employment",
        "service contract",
    }
)


def _as_bool(value: Any, default: bool = False) -> bool:
    if value is None:
        return default
    if isinstance(value, bool):
        return value
    if isinstance(value, (int, float)):
        return bool(value)
    text = str(value).strip().lower()
    if text in {"true", "yes", "1", "y"}:
        return True
    if text in {"false", "no", "0", "n", "null", "none", ""}:
        return False
    return default


def _normalize_detected_kind(raw: Any) -> str | None:
    text = str(raw or "").strip().lower().replace("-", "_")
    text = re.sub(r"\s+", " ", text)
    compact = text.replace(" ", "_")
    if text in _APPOINTMENT_KINDS or compact in _APPOINTMENT_KINDS:
        return "appointment"
    if text in _CONTRACT_KINDS or compact in _CONTRACT_KINDS:
        return "contract"
    if not text or text in {"other", "unknown", "none", "n/a", "na"}:
        return None
    return "other"


def _letter_label(kind: str) -> str:
    return "appointment letter" if kind == "appointment" else "employment contract"


def _docx_text_excerpt(content: bytes, max_chars: int = 2800) -> str:
    try:
        from docx import Document
    except ImportError:
        return ""
    try:
        doc = Document(BytesIO(content))
        parts: list[str] = []
        for para in doc.paragraphs:
            text = (para.text or "").strip()
            if text:
                parts.append(text)
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if (c.text or "").strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)[:max_chars]
    except Exception:  # noqa: BLE001
        logger.exception("Could not extract text from stored letter for verification")
        return ""


def _issued_letter_excerpt(db: Session, employee_id: int, kind: str) -> str:
    doc = _latest_letter(db, employee_id, kind)
    if doc is None:
        return ""
    try:
        content = read_stored_file(doc.file_path)
    except Exception:  # noqa: BLE001
        logger.exception("Could not read stored letter for verification context")
        return ""
    return _docx_text_excerpt(content)


def decide_letter_verification(kind: str, vision: dict[str, Any]) -> tuple[bool, str, str]:
    """Map vision JSON to (verified, status, message). Pure — used by tests."""
    expected = _letter_label(kind)
    other = _letter_label("contract" if kind == "appointment" else "appointment")
    notes = str(vision.get("notes") or "").strip() or None
    readable = _as_bool(
        vision.get("image_readable", vision.get("readable")),
        default=False,
    )
    looks_like = _as_bool(
        vision.get(
            "looks_like_expected_letter",
            vision.get("looks_like_letter_document"),
        ),
        default=False,
    )
    has_sig = _as_bool(
        vision.get(
            "has_handwritten_signature",
            vision.get("has_client_signature"),
        ),
        default=False,
    )
    detected = _normalize_detected_kind(
        vision.get("detected_document_kind", vision.get("document_kind"))
    )
    wrong_type = detected in {"appointment", "contract"} and detected != kind
    correct_type = (not wrong_type) and (looks_like or detected == kind)

    if not readable:
        return (
            False,
            "unreadable",
            notes
            or "The file is too unclear to identify the letter. Upload a sharper PDF or photo of the signed page.",
        )
    if wrong_type:
        return (
            False,
            "wrong_type",
            notes
            or f"This looks like {other}, not the {expected}. Upload the correct signed document.",
        )
    if not correct_type:
        return (
            False,
            "not_letter",
            notes
            or f"This file does not look like the {expected}. Upload a clear PDF or photo of the signed {expected}.",
        )
    if not has_sig:
        return (
            False,
            "no_signature",
            notes
            or f"This is the {expected}, but no handwritten signature was found. Ask them to sign, then upload again.",
        )
    return True, "verified", notes or "Signature found — letter marked Verified."


def _parse_vision_json(text: str) -> dict[str, Any]:
    raw = (text or "").strip()
    if raw.startswith("```"):
        raw = re.sub(r"^```(?:json)?\s*", "", raw)
        raw = re.sub(r"\s*```$", "", raw)
    try:
        data = json.loads(raw)
        if isinstance(data, dict):
            return data
    except json.JSONDecodeError:
        pass
    match = re.search(r"\{[\s\S]*\}", raw)
    if match:
        try:
            data = json.loads(match.group(0))
            if isinstance(data, dict):
                return data
        except json.JSONDecodeError:
            pass
    raise ValidationFailed("Could not parse signature verification response from AI")


def _ai_check_signature(
    *,
    image_bytes: bytes,
    mime_type: str,
    kind: str,
    employee_name: str,
    employee_code: str = "",
    employee_cnic: str = "",
    role_title: str = "",
    issued_letter_excerpt: str = "",
) -> dict[str, Any]:
    settings = get_settings()
    api_keys = settings.resolved_gemini_letter_api_keys()
    if not api_keys:
        raise ValidationFailed(
            "Letter signature verification requires GEMINI_LETTER_API_KEY or GEMINI_API_KEY "
            "on the backend."
        )
    expected = _letter_label(kind)
    other = _letter_label("contract" if kind == "appointment" else "appointment")
    excerpt_block = (
        f"\nIssued {expected} text excerpt (for matching, not as proof of signature):\n"
        f"---\n{issued_letter_excerpt.strip()}\n---\n"
        if issued_letter_excerpt.strip()
        else ""
    )
    prompt = f"""You are verifying a signed HR document (photo or PDF scan) for Kafi Group.

Task: decide whether this file is the employee's {expected} AND whether a handwritten signature is on it.

Employee on file:
- full_name: {employee_name or "unknown"}
- employee_code: {employee_code or "unknown"}
- cnic: {employee_cnic or "unknown"}
- role_title: {role_title or "unknown"}
- expected_document: {expected}
{excerpt_block}
Return ONLY valid JSON (no markdown) with this exact shape:
{{
  "image_readable": true,
  "detected_document_kind": "appointment",
  "looks_like_expected_letter": true,
  "has_handwritten_signature": true,
  "employee_name_visible": true,
  "visible_title": "string or null",
  "notes": "one short sentence explaining the decision"
}}

detected_document_kind must be one of: "appointment", "contract", "other".

What counts as the {expected}:
- A printed, photographed, or PDF-scanned Kafi Group {expected} (letterhead, title, appointment/contract wording, employee name, role, CNIC, salary/terms, signature block).
- A last/signature page of that same letter is valid if it is clearly the acknowledgement/acceptance/signature page of the {expected} (not a random scrap).
- Match the employee name/CNIC/role when they are readable.

What does NOT count:
- CNIC cards, education documents, selfies, invoices, unrelated letters, blank paper, screenshots of chat.
- The other HR letter type ({other}) — set detected_document_kind to that type and looks_like_expected_letter=false.

Signature rules:
- has_handwritten_signature=true ONLY if you can see a human handwritten / ink / wet signature, initials, or signed mark in a signature area.
- Printed names, typed "/sd", stamps alone, or a blank signature line do NOT count.
- If the page is the right letter but unsigned, looks_like_expected_letter=true and has_handwritten_signature=false.

If the file is too blurry, dark, or cropped to judge, set image_readable=false and has_handwritten_signature=false.
"""
    response = generate_content_with_fallback(
        api_keys=api_keys,
        models=settings.resolved_gemini_letter_models(),
        prompt=[
            prompt,
            {"mime_type": mime_type or "image/jpeg", "data": image_bytes},
        ],
        pool_id="letter_verify",
    )
    return _parse_vision_json(getattr(response, "text", "") or "")


def verify_letter_signature(
    db: Session,
    auth: AuthContext,
    employee_id: int,
    kind: str,
    *,
    filename: str,
    content: bytes,
    mime_type: str | None,
) -> LetterSignatureVerifyResult:
    employee = employee_service.get_employee(db, employee_id)
    if kind not in LETTER_CATEGORIES:
        raise ValidationFailed(f"Unknown letter type '{kind}'")
    if _latest_letter(db, employee_id, kind) is None:
        raise BusinessRuleViolation(_MISSING)

    resolved_mime = _verify_upload_mime(filename or "", mime_type)
    if not content:
        raise ValidationFailed("Uploaded file is empty")

    excerpt = _issued_letter_excerpt(db, employee_id, kind)
    try:
        vision = _ai_check_signature(
            image_bytes=content,
            mime_type=resolved_mime,
            kind=kind,
            employee_name=(employee.full_name or "").strip(),
            employee_code=(employee.employee_code or "").strip(),
            employee_cnic=(employee.cnic or "").strip(),
            role_title=(employee.role_title or "").strip(),
            issued_letter_excerpt=excerpt,
        )
    except ValidationFailed:
        raise
    except Exception as exc:  # noqa: BLE001
        logger.exception("Letter signature AI failed")
        raise ValidationFailed(f"AI signature check failed: {exc}") from exc

    verified, status, message = decide_letter_verification(kind, vision)
    if not verified:
        return LetterSignatureVerifyResult(
            verified=False,
            status=status,
            message=message,
            kind=kind,
            employee_id=employee.id,
        )

    signed_cat = _signed_category(kind)
    _delete_category_docs(db, employee.id, signed_cat)
    path, stored_mime = store_employee_file(
        employee_id=employee.id,
        filename=filename or ("signed_letter.pdf" if resolved_mime == "application/pdf" else "signed_letter.jpg"),
        content=content,
        subdir="letters",
        images_only=False,
    )
    row = EmployeeDocument(
        employee_id=employee.id,
        category=signed_cat,
        title=(
            "Appointment letter — signed (verified)"
            if kind == "appointment"
            else "Employment contract — signed (verified)"
        ),
        file_path=path,
        original_filename=filename or ("signed_letter.pdf" if resolved_mime == "application/pdf" else "signed_letter.jpg"),
        mime_type=stored_mime or resolved_mime,
    )
    db.add(row)
    db.flush()
    notes = str(vision.get("notes") or "").strip() or None
    audit_service.log_from_auth(
        db,
        auth,
        action=f"employee.letter_{kind}_verified",
        entity_type="employee",
        entity_id=employee.id,
        after_state={
            "kind": kind,
            "document_id": row.id,
            "filename": row.original_filename,
            "ai_notes": notes,
            "detected_document_kind": vision.get("detected_document_kind"),
        },
    )
    return LetterSignatureVerifyResult(
        verified=True,
        status="verified",
        message=message,
        kind=kind,
        employee_id=employee.id,
        document_id=row.id,
    )
