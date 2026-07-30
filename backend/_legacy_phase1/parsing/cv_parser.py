"""Extracts plain text from CV files (PDF / DOCX / TXT) so it can be fed to
the scoring LLM."""
from __future__ import annotations

from pathlib import Path

import pdfplumber
from docx import Document


class UnsupportedCVFormat(Exception):
    pass


def extract_text(file_path: str | Path) -> str:
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return _extract_pdf(path)
    if suffix == ".docx":
        return _extract_docx(path)
    if suffix == ".txt":
        return path.read_text(encoding="utf-8", errors="ignore")

    raise UnsupportedCVFormat(
        f"Unsupported CV file type '{suffix}' for {path.name}. Supported: .pdf, .docx, .txt"
    )


def _extract_pdf(path: Path) -> str:
    chunks: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            chunks.append(text)
    return "\n".join(chunks).strip()


def _extract_docx(path: Path) -> str:
    document = Document(path)
    paragraphs = [p.text for p in document.paragraphs]

    for table in document.tables:
        for row in table.rows:
            paragraphs.append(" | ".join(cell.text for cell in row.cells))

    return "\n".join(paragraphs).strip()
